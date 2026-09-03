import io

import docx
import PyPDF2
import pytesseract
from PIL import Image


def extract_uploaded_text(uploaded_file):
    extension = uploaded_file.name.rsplit(".", 1)[-1].lower()
    data = uploaded_file.getvalue()
    if extension == "txt":
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")
    if extension == "pdf":
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    if extension == "docx":
        document = docx.Document(io.BytesIO(data))
        parts = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n\n".join(parts)
    if extension in {"png", "jpg", "jpeg"}:
        image = Image.open(io.BytesIO(data))
        return pytesseract.image_to_string(image, lang="eng+chi_sim")
    raise ValueError(f"不支持的文件类型：{extension}")
